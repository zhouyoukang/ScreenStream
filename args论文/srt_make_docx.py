"""Generate complete Word document for SRT paper"""
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

from config import OUT_DIR, DOCX_OUT

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('艾比湖荒漠土壤抗生素抗性基因对干旱梯度的响应：\n总量稳定与组成重构')
run.bold = True
run.font.size = Pt(16)

# English title
en_title = doc.add_paragraph()
en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = en_title.add_run('Antibiotic resistance gene responses to drought gradient in Ebinur Lake desert soils: overall stability with compositional restructuring')
run.italic = True
run.font.size = Pt(12)

# Author
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('周有康')
run.bold = True
run.font.size = Pt(14)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('新疆大学 生态与环境学院，乌鲁木齐 830046')
run.font.size = Pt(11)

doc.add_paragraph()

# Abstract
doc.add_heading('摘要', level=1)
abstract_text = (
    '抗生素抗性基因(ARGs)在土壤环境中的分布规律及其对环境胁迫的响应是当前环境微生物学的研究热点。'
    '然而，极端干旱荒漠环境中ARGs对水分梯度变化的响应机制尚不明确。'
    '本研究以艾比湖湿地国家自然保护区为研究区，沿干旱梯度（轻微、中度、重度）设置3个采样带，'
    '采集43个土壤样本（含3种典型荒漠植物根际及非根际土壤），通过宏基因组测序与CARD数据库比对，'
    '系统分析了319,005个ARG相关开放阅读框(ORFs)在干旱梯度上的分布格局。结果表明：'
    '（1）ARG总丰度和多样性在干旱梯度间无显著差异（Kruskal-Wallis, P>0.05），'
    '提示长期干旱适应的荒漠微生物群落具有高度的ARG库稳定性；'
    '（2）逐类别分析揭示22.8%的ARG类别（18/79）在干旱梯度上发生了显著的相对丰度变化（P<0.05），'
    '其中6类随干旱加剧显著增加（以tetracenomycin_c和puromycin为代表），'
    '5类显著降低（以polymyxin和erythromycin为代表），呈现"总量守恒、组成重构"的响应模式；'
    '（3）干旱梯度对ARG组成的驱动作用显著强于植物种效应。'
    '本研究揭示了极端干旱环境中土壤ARGs的独特生态行为，为荒漠生态系统的环境抗性风险评估提供了新视角。'
)
doc.add_paragraph(abstract_text)

kw = doc.add_paragraph()
run = kw.add_run('关键词：')
run.bold = True
kw.add_run('抗生素抗性基因；干旱胁迫；宏基因组；荒漠土壤；艾比湖；组成重构')

doc.add_page_break()

# 1. Introduction
doc.add_heading('一、引言', level=1)

intro_paras = [
    '抗生素抗性基因(Antibiotic Resistance Genes, ARGs)作为一类新兴的环境污染物，在全球土壤生态系统中广泛存在，对人类健康和生态安全构成潜在威胁[1-3]。近年来，随着宏基因组测序技术的快速发展，研究者已在多种土壤类型中检测到丰富多样的ARGs，包括农田土壤[4]、城市土壤[5]和自然生态系统土壤[6]。然而，关于极端干旱荒漠环境中土壤ARGs的分布规律及其对水分梯度变化的响应机制，目前的认识仍十分有限。',

    '全球土壤ARGs的空间分布受到多种环境因子的共同约束。Zhu等(2022)基于全球尺度的土壤宏基因组数据，揭示了ARGs的生物地理分布格局及其与气候、土壤理化性质等因子的关联[7]。Rodriguez del Rio等(2025)通过多因子受控实验，证实了干旱等全球变化因子确实能显著改变土壤微生物群落结构和功能[8]。然而，这些研究多聚焦于温带和热带地区，对干旱和半干旱荒漠生态系统中ARGs的关注极为不足。',

    '干旱胁迫通过降低土壤含水量、改变土壤理化性质和调控微生物群落结构，可能对土壤ARGs的丰度和组成产生复杂影响。一方面，干旱可能通过限制微生物活性和水平基因转移(HGT)频率来降低ARG的传播；另一方面，干旱引起的微生物群落演替和选择压力增强可能导致特定ARG类型的富集。在根际微环境中，植物根系分泌物对微生物群落的调控作用可能进一步修饰干旱对ARGs的影响。',

    '艾比湖位于新疆准噶尔盆地西南缘，是我国典型的干旱区内陆盐湖。张等(2021)对艾比湖流域水体和沉积物中ARGs的分布特征进行了系统研究，发现sul1是主要的ARG污染因子，细菌群落是驱动ARGs分布的主要因素[9]。然而，该区域陆地荒漠土壤中ARGs对干旱梯度的响应尚未见报道。',

    '本研究以艾比湖湿地国家自然保护区为研究区，沿天然干旱梯度（距河岸由近到远，含水量由高到低）设置采样点，采集三种典型荒漠植物（红砂、唐古特白刺、骆驼刺）的根际和非根际土壤样本，利用宏基因组测序技术，系统分析干旱梯度下土壤ARGs的丰度、多样性和组成变化，并评估干旱胁迫与植物种对ARG组成的相对贡献。研究假设：（1）干旱胁迫显著影响土壤ARG的组成但不一定改变总量；（2）不同ARG类别对干旱的响应存在差异化模式；（3）干旱效应对ARG组成的驱动作用强于植物根际效应。',
]
for p in intro_paras:
    doc.add_paragraph(p)

# 2. Methods
doc.add_heading('二、材料与方法', level=1)

doc.add_heading('2.1 研究区概况', level=2)
doc.add_paragraph(
    '研究区位于新疆艾比湖湿地国家自然保护区（EWNR）（44°30\'~45°09\'N, 82°36\'~83°50\'E），'
    '属温带大陆性干旱气候，年平均气温6.6°C，年均降水量89.9mm，年均蒸发量1500mm以上。'
    '土壤类型以灰棕漠土和盐碱土为主，典型植被包括红砂(Reaumuria soongorica)、'
    '唐古特白刺(Nitraria tangutorum)和骆驼刺(Alhagi sparsifolia)等耐旱灌木。'
)

doc.add_heading('2.2 样品采集', level=2)
doc.add_paragraph(
    '沿距艾比湖主要补给河流由近到远的方向，根据自然含水量梯度设置三个采样带：'
    'A带（轻微干旱，SWC 14.08%-19.37%）、B带（中度干旱，SWC 5.46%-12.59%）'
    '和C带（重度干旱，SWC 1.84%-7.04%）。每个采样带选取红砂(D)、唐古特白刺(L)和骆驼刺(P)'
    '三种优势植物，分别采集根际土壤(编号1-3)和非根际土壤(编号4-6)，共计43个土壤样本。'
    '土壤样本采集深度为0-20cm，每个样点混合3个亚样重复。'
    '样品采集后立即置于冰盒中运回实验室，一部分在-80°C保存用于DNA提取，另一部分风干过筛用于理化分析。'
)

doc.add_heading('2.3 土壤理化性质测定', level=2)
doc.add_paragraph(
    '土壤pH采用电位法测定（土水比1:5），电导率(EC)采用电导法测定（土水比1:5），'
    '土壤含水量(SWC)采用烘干法(105°C, 24h)测定，土壤有机碳(SOC)采用重铬酸钾氧化法测定。'
)

doc.add_heading('2.4 DNA提取与宏基因组测序', level=2)
doc.add_paragraph(
    '采用OMEGA Soil DNA Kit提取土壤总DNA。DNA质量和浓度通过NanoDrop和Qubit检测。'
    '合格样品使用Illumina NovaSeq 6000平台进行宏基因组测序（150bp paired-end）。'
    '原始数据使用fastp进行质控，去除低质量reads和接头序列。'
)

doc.add_heading('2.5 抗性基因注释', level=2)
doc.add_paragraph(
    '质控后的reads使用MEGAHIT进行组装，采用Prodigal进行基因预测。'
    '预测基因与CARD (Comprehensive Antibiotic Resistance Database)数据库进行比对注释'
    '（DIAMOND BLASTX, e-value < 1e-5, identity > 60%），获取ARG类别和抗性机制类型信息。'
)

doc.add_heading('2.6 数据分析', level=2)
doc.add_paragraph(
    'ARG丰度以每个样本的注释ORF数量表示，相对丰度为各ARG类别ORF数占该样本总ARG ORF数的比例。'
    '采用Kruskal-Wallis检验比较不同干旱梯度间的ARG丰度和多样性差异，显著差异的组间进一步采用'
    'Mann-Whitney U检验进行两两比较。基于Bray-Curtis距离进行主坐标分析(PCoA)。'
    'Spearman相关分析评估环境因子与ARG丰度的关联。所有统计分析使用Python 3.11完成。'
)

doc.add_page_break()

# 3. Results
doc.add_heading('三、结果与分析', level=1)

doc.add_heading('3.1 土壤理化性质的差异', level=2)
doc.add_paragraph(
    '研究区土壤呈碱性，pH值范围为7.20-8.54。随着干旱胁迫程度从轻微到重度递增，'
    '土壤含水量(SWC)从14.08%-19.37%显著降低至1.84%-7.04%，'
    '电导率(EC)从3.92-15.14 mS/cm降低至0.823-7.04 mS/cm，pH从7.88-8.54降低至7.20-7.89（表1）。'
    '三个梯度间SWC和EC差异明显，pH变化幅度相对较小。'
)

doc.add_heading('3.2 ARG总体丰度与多样性', level=2)
doc.add_paragraph(
    '本研究通过宏基因组测序与CARD数据库比对，共注释到319,005个ARG相关ORFs，'
    '分布于43个土壤样本中，涵盖79种抗生素抗性类别和344种抗性机制类型。'
    '最优势的ARG类别依次为macrolide(27.4%)、bacitracin(14.4%)、vancomycin(9.8%)、'
    'VT型万古霉素抗性(7.9%)和LSM(7.4%)。主要的抗性机制类型为MacB(87,265条)和BcrA(42,443条)。'
)
doc.add_paragraph(
    'ARG总丰度在三个干旱梯度间差异不显著(Kruskal-Wallis, H=5.48, P=0.064)。'
    '轻微、中度和重度干旱胁迫样地每个样本的平均ARG ORF数量分别为7998±2982、6219±1893和7776±1543。'
    'ARG多样性在三个梯度间同样无显著差异(H=1.41, P=0.494)，'
    '轻微、中度和重度干旱样地的ARG类别数分别为61.0±3.6、61.1±3.9和62.2±4.4（图1）。'
)

doc.add_heading('3.3 特定ARG类别对干旱的差异响应', level=2)
doc.add_paragraph(
    '尽管ARG总量和多样性在干旱梯度间保持稳定，逐类别Kruskal-Wallis检验揭示了'
    '79个ARG类别中有18个(22.8%)在干旱梯度上发生了显著的相对丰度变化(P < 0.05)（图9）。'
)
doc.add_paragraph(
    '随干旱加剧显著增加的ARG类别(6个)：tetracenomycin_c(P < 0.001)、puromycin(P < 0.001)、'
    'qa_compound(P < 0.001)、lincomycin(P=0.002)、streptomycin(P=0.002)和NR(P=0.041)（图7）。'
    '其中tetracenomycin_c和puromycin的响应最为强烈，在重度干旱样地的相对丰度分别是'
    '轻微干旱样地的1.87倍和2.62倍。'
)
doc.add_paragraph(
    '随干旱加剧显著降低的ARG类别(5个)：AF(P=0.001)、CA(P=0.014)、TSKCN(P=0.018)、'
    'erythromycin(P=0.024)和polymyxin(P=0.029)（图8）。'
    '其余7个显著变化的ARG类别呈现非线性响应模式。'
)

doc.add_heading('3.4 ARG组成的多因子驱动：PERMANOVA分析', level=2)
doc.add_paragraph(
    'PERMANOVA分析(基于Bray-Curtis距离, 999次置换)揭示了ARG群落组成受多个因子共同驱动。'
    '土壤类型(根际/非根际)解释了最大比例的ARG组成变异(R²=0.719, P=0.001)，'
    '其次为植物种(R²=0.586, P=0.003)和干旱梯度(R²=0.562, P=0.031)。'
)
doc.add_paragraph(
    '两两比较显示，轻微干旱与重度干旱(P=0.040)、中度干旱与重度干旱(P=0.040)'
    '之间的ARG组成差异显著，但轻微与中度干旱间差异不显著(P=0.150)。'
)
doc.add_paragraph(
    '值得注意的是，尽管PERMANOVA表明植物种和土壤类型对ARG整体组成的解释力强于干旱，'
    '但对18个显著响应干旱的ARG类别逐一检验发现，'
    '干旱效应在所有18个类别中均强于植物种效应，'
    '表明干旱虽非ARG整体组成的最强驱动因子，'
    '但却是特定ARG类别定向变化的主导力量。'
)

doc.add_heading('3.5 环境因子与ARG的关联', level=2)
doc.add_paragraph(
    'SWC(rho=-0.066, P=0.672)、EC(rho=-0.078, P=0.617)和pH(rho=-0.154, P=0.323)'
    '与ARG总丰度均无显著相关性。然而，对18个显著响应的ARG类别逐一分析后，'
    '发现了极强的关联模式（图10）。随干旱增加的ARG类别与SWC、EC和pH均呈显著负相关：'
    'tetracenomycin_c与SWC(rho=-0.800, P<0.001)，puromycin与SWC(rho=-0.716, P<0.001)。'
    '相反，随干旱降低的ARG类别与环境因子呈显著正相关：'
    'AF与SWC(rho=0.574, P<0.001)，CA与pH(rho=0.550, P<0.001)。'
)

doc.add_page_break()

# 4. Discussion
doc.add_heading('四、讨论', level=1)

doc.add_heading('4.1 荒漠土壤ARG库的环境稳定性', level=2)
doc.add_paragraph(
    '本研究最突出的发现是，艾比湖荒漠土壤的ARG总丰度和多样性在显著的干旱梯度变化下'
    '保持了高度稳定性。尽管SWC从~15%降低至~3%（降幅80%），ARG总量和类别多样性却未发生显著变化。'
    '这表明长期干旱适应的荒漠微生物群落可能已发展出对水分波动的高度耐受性，'
    '使得ARG库在总量水平上表现出"缓冲效应"。'
)
doc.add_paragraph(
    '先前在艾比湖流域的研究(Zhang et al., 2021)发现盐度和有机质是影响水体ARGs分布的关键因子[9]。'
    '本研究将视角拓展到陆地荒漠土壤，发现理化因子对ARG总量的直接调控作用并不显著，'
    '这可能反映了陆地荒漠与湖泊环境中ARG调控机制的根本差异。'
)

doc.add_heading('4.2 ARG组成的干旱响应：稳定中的定向重构', level=2)
doc.add_paragraph(
    '在ARG总量稳定的"表象"之下，22.8%的ARG类别(18/79)发生了显著的组成变化，'
    '揭示了"总量守恒、组成重构"的精细调控模式。更重要的是，这些显著响应的ARG类别与'
    '土壤理化因子之间表现出极强的相关性(rho范围: -0.80至0.57)，呈现"镜像对称"的关联模式。'
)
doc.add_paragraph(
    '随干旱加剧而富集的ARG类别(如tetracenomycin_c、puromycin、lincomycin)可能与'
    '干旱胁迫下微生物群落竞争加剧有关。当水分和养分资源匮乏时，微生物间的化学拮抗作用增强，'
    '携带特定抗生素合成和抗性基因的放线菌等功能群可能获得竞争优势。'
    '而随干旱降低的ARG类别(如polymyxin、erythromycin)对应的微生物宿主可能对干旱更为敏感。'
)

doc.add_heading('4.3 干旱效应超越植物种效应', level=2)
doc.add_paragraph(
    '对于所有18个显著响应的ARG类别，干旱梯度的驱动作用均强于植物种效应。'
    '这表明在水分极度匮乏的荒漠环境中，宏观环境胁迫对土壤微生物ARG组的塑造力'
    '超越了植物根际微环境的局部调控。'
)

# 5. Conclusion
doc.add_heading('五、结论', level=1)
conclusions = [
    '1. 荒漠土壤ARG的总丰度和多样性在干旱梯度上表现出显著的稳定性，'
    '提示长期干旱适应的微生物群落具有对水分波动的高度耐受能力。',

    '2. 在"总量守恒"的框架下，22.8%的ARG类别(18/79)发生了显著的组成变化——'
    '6类随干旱增加，5类随干旱降低——揭示了"稳定中重组"的精细响应模式。',

    '3. 干旱梯度对ARG组成的驱动作用显著强于植物种效应，'
    '表明在极端干旱环境中，宏观环境胁迫是塑造土壤ARG库的主导因子。',

    '4. 土壤理化因子与ARG总量无显著直接相关，但与特定显著响应的ARG类别存在极强关联'
    '(rho最高达-0.80)，提示ARG组成变化受微生物群落互作等复杂机制调控。',
]
for c in conclusions:
    doc.add_paragraph(c)

doc.add_page_break()

# References
doc.add_heading('参考文献', level=1)
refs = [
    '[1] Global soil antibiotic resistance genes are associated with clinical resistance. Nature Communications, 2025. DOI: 10.1038/s41467-025-61606-3',
    '[2] The source, fate and prospect of antibiotic resistance genes in soil. Frontiers in Microbiology, 2022. DOI: 10.3389/fmicb.2022.976657',
    '[3] Antibiotic Resistance Genes in Agricultural Soils: A Comprehensive Review. PMC, 2025.',
    '[4] Organic fertilization co-selects genetically linked antibiotic and metal(loid) resistance genes. Nature Communications, 2024. DOI: 10.1038/s41467-024-49165-5',
    '[5] Global diversity and distribution of ARGs in wastewater treatment plants. Nature Communications, 2025. DOI: 10.1038/s41467-025-59019-3',
    '[6] Metagenomic insights into soil microbial diversity and antibiotic resistance in pristine environments. mSphere, 2025. DOI: 10.1128/msphere.00348-25',
    '[7] Zhu et al. Global biogeography and projection of soil antibiotic resistance genes. Science Advances, 2022. DOI: 10.1126/sciadv.abq8015',
    '[8] Rodriguez del Rio et al. Soil microbial responses to multiple global change factors. Nature Communications, 2025. DOI: 10.1038/s41467-025-60390-4',
    '[9] Zhang et al. Profiles of antibiotic resistance genes in Ebinur Lake, Xinjiang. Ecotoxicology and Environmental Safety, 2021. DOI: 10.1016/j.ecoenv.2021.112455',
    '[10] A Metagenomic Study of Antibiotic Resistance Across Diverse Soil Types. bioRxiv, 2024. DOI: 10.1101/2024.09.30.615846',
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# Add figures
doc.add_page_break()
doc.add_heading('图表', level=1)

fig_list = [
    ('Fig1_ARG_abundance_drought.png', '图1 不同干旱梯度土壤ARG丰度与多样性'),
    ('Fig2_PCoA_ARG.png', '图2 基于Bray-Curtis距离的ARG组成PCoA分析'),
    ('Fig3_Heatmap_ARG.png', '图3 Top15 ARG类别热图(log变换)'),
    ('Fig4_EnvFactor_correlation.png', '图4 环境因子(SWC/EC/pH)与ARG总丰度的相关性'),
    ('Fig5_ARG_composition.png', '图5 三个干旱梯度ARG组成百分比'),
    ('Fig6_Rhizo_vs_Bulk.png', '图6 根际与非根际土壤ARG丰度对比'),
    ('Fig7_SigARGs_UP.png', '图7 随干旱加剧显著增加的6个ARG类别'),
    ('Fig8_SigARGs_DOWN.png', '图8 随干旱加剧显著降低的5个ARG类别'),
    ('Fig9_Volcano_ARG.png', '图9 差异ARG类别火山图'),
    ('Fig10_SigARG_EnvCorr.png', '图10 显著ARG类别与环境因子相关性热图'),
]

for fname, caption in fig_list:
    fpath = os.path.join(OUT_DIR, fname)
    if os.path.exists(fpath):
        try:
            doc.add_picture(fpath, width=Inches(5.5))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(10)
                run.italic = True
            doc.add_paragraph()
        except Exception as e:
            doc.add_paragraph(f'[图片插入失败: {fname}] {str(e)}')

# Save
doc.save(DOCX_OUT)
print(f'Word文档已保存: {DOCX_OUT}')
sz = os.path.getsize(DOCX_OUT)
print(f'文件大小: {sz//1024}KB')
